import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Guia_Comentarios_Workiva.docx"

AZUL_OSCURO   = RGBColor(0x1F, 0x49, 0x7D)
AZUL_CLARO_BG = "BDD7EE"
AMARILLO_BG   = "FFFF99"
ROJO          = RGBColor(0xC0, 0x00, 0x00)
GRIS_CAPTION  = RGBColor(0x80, 0x80, 0x80)
FONT = "Calibri"


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name   = FONT
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_top_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "1F497D")
    pBdr.append(top)
    pPr.append(pBdr)


def add_image_or_placeholder(doc, img_path, caption_text, width_inches=4.5):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        p = doc.add_paragraph("[ IMAGEN: " + os.path.basename(img_path) + " ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]
        set_font(r, size=10, italic=True, color=GRIS_CAPTION)
    cap = doc.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.runs[0]
    set_font(r, size=9, italic=True, color=GRIS_CAPTION)
    doc.add_paragraph()


def add_step_header(doc, step_text):
    p = doc.add_paragraph()
    add_top_border(p)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(step_text)
    set_font(r, size=13, bold=True, color=AZUL_OSCURO)
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size)
    return p


# ── Documento ─────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Titulo
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titulo.add_run("Guia de comentarios con permisos personalizados en Workiva")
set_font(r, size=18, bold=True, color=AZUL_OSCURO)
doc.add_paragraph()

# Banner
banner = doc.add_paragraph()
banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
banner.paragraph_format.left_indent  = Cm(0.5)
banner.paragraph_format.right_indent = Cm(0.5)
banner.paragraph_format.space_before = Pt(6)
banner.paragraph_format.space_after  = Pt(6)
r = banner.add_run(
    "A partir de ahora, todos los comentarios en Workiva deben crearse con "
    "permisos personalizados. Esta guia explica como hacerlo correctamente."
)
set_font(r, size=11, bold=True)
shade_paragraph(banner, AZUL_CLARO_BG)
doc.add_paragraph()

# ── Paso 1 ────────────────────────────────────────────────────────────────────
add_step_header(doc, "Paso 1 — Abrir el panel de comentarios")
add_body(doc,
    "En el documento o spreadsheet de Wdesk, haz clic en el icono de comentarios "
    "en la barra lateral derecha. Se abrira el panel con el campo \"New comment...\"")

add_image_or_placeholder(
    doc, "img1.png",
    "Imagen 1: Panel de comentarios — escribiendo la @mencion en el campo de texto."
)

# ── Paso 2 ────────────────────────────────────────────────────────────────────
add_step_header(doc, "Paso 2 — Redactar el comentario y mencionar a quien corresponda")
add_body(doc,
    "Escribe el comentario en el campo de texto. Si necesitas notificar a alguien "
    "especificamente, usa @nombre para mencionarlo (ej: @EmersonGarridoVillarroel).")

# ── Paso 3 ────────────────────────────────────────────────────────────────────
add_step_header(doc, "Paso 3 — ⚠ PASO CLAVE: Ampliar la seccion de permisos antes de publicar")
add_body(doc,
    "No hagas clic directamente en \"Post\". En cambio, haz clic en la flecha pequena (▲) "
    "que aparece al lado derecho del boton verde Post. Esto despliega el panel de "
    "\"Comment Permissions\".")

adv = doc.add_paragraph()
r1 = adv.add_run("ADVERTENCIA: ")
set_font(r1, size=11, bold=True)
r2 = adv.add_run(
    "Si publicas directamente con \"Post\" sin configurar permisos, el comentario "
    "queda visible para todos los usuarios del archivo, sin restricciones.")
set_font(r2, size=11)
shade_paragraph(adv, AMARILLO_BG)

add_image_or_placeholder(
    doc, "img2.png",
    "Imagen 2: Flecha ▲ junto al boton Post desplegando el panel Comment Permissions."
)

# ── Paso 4 ────────────────────────────────────────────────────────────────────
add_step_header(doc, "Paso 4 — Seleccionar quienes pueden ver y resolver el comentario")
add_body(doc,
    "En el panel \"Comment Permissions\", usa el buscador para encontrar personas o grupos. "
    "Ejemplo de seleccion correcta:")

for item in ["Estados Financieros (grupo)", "Emerson Garrido Villarroel (persona)"]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item)
    set_font(r, size=11)

add_image_or_placeholder(
    doc, "img3.png",
    "Imagen 3: Permisos seleccionados — Estados Financieros y Emerson Garrido Villarroel."
)

# ── Paso 5 ────────────────────────────────────────────────────────────────────
add_step_header(doc, "Paso 5 — Publicar el comentario")
add_body(doc,
    "Una vez seleccionados los permisos, haz clic en Post (aparecera con icono de "
    "llave 🔑, confirmando que los permisos estan activos). El comentario quedara "
    "visible unicamente para las personas seleccionadas.")

add_image_or_placeholder(
    doc, "img4.png",
    "Imagen 4: Boton Post con icono de llave 🔑 activo, confirmando permisos configurados."
)

# ── Resultado esperado ────────────────────────────────────────────────────────
doc.add_paragraph()
res_hdr = doc.add_paragraph()
add_top_border(res_hdr)
r = res_hdr.add_run("Resultado esperado")
set_font(r, size=13, bold=True, color=AZUL_OSCURO)

checks = [
    "El comentario aparece con el nombre del autor y la hora de publicacion.",
    "Muestra en \"Comment Permissions\" solo las personas/grupos autorizados.",
    "Hay un campo \"Reply or mention others with @\" para respuestas en hilo.",
    "El boton de opciones (▲) permite editar permisos o acceder a otras acciones.",
]
for c in checks:
    p = doc.add_paragraph()
    r = p.add_run("✓  " + c)
    set_font(r, size=11)

add_image_or_placeholder(
    doc, "img5.png",
    "Imagen 5: Comentario publicado correctamente con permisos aplicados."
)

# ── Regla Fundamental ─────────────────────────────────────────────────────────
doc.add_paragraph()
regla_hdr = doc.add_paragraph()
add_top_border(regla_hdr)
r = regla_hdr.add_run("⛔  Regla Fundamental")
set_font(r, size=13, bold=True, color=ROJO)

regla_body = doc.add_paragraph()
r1 = regla_body.add_run("Los comentarios se RESUELVEN, no se eliminan.\n")
set_font(r1, size=11, bold=True, color=ROJO)
r2 = regla_body.add_run(
    "Al hacer clic en el menu de opciones (▲) de un comentario, siempre usa "
    "\"Resolve comment\" para cerrar un comentario una vez atendido.\n"
    "Eliminar borra el historial; resolver lo archiva y mantiene la trazabilidad.")
set_font(r2, size=11, bold=False, color=ROJO)

add_image_or_placeholder(
    doc, "img6.png",
    "Imagen 6: Menu de opciones mostrando \"Resolve comment\" como accion correcta."
)

doc.save(OUT)
print("Documento generado:", OUT)
